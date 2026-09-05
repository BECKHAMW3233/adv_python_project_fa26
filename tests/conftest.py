"""Shared pytest fixture builders -- small MOS/training/program files built at test time
rather than relying on the full supplied source files, per the assignment spec."""

from __future__ import annotations

from pathlib import Path

import docx
import openpyxl


def build_mos_workbook(path: Path, sheets: dict[str, list[tuple]]) -> None:
    """sheets: {sheet_name: [row_tuples, ...]} -- each row written verbatim, one cell per tuple item."""
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)
    for sheet_name, rows in sheets.items():
        worksheet = workbook.create_sheet(sheet_name)
        for row in rows:
            worksheet.append(row)
    workbook.save(path)


def build_training_docx(path: Path, tables: list[list[list[str]]], narrative: str = "") -> None:
    """tables: list of tables, each a list of rows, each row a list of cell strings.
    narrative: optional body-paragraph text (outside any table) to confirm it's ignored."""
    document = docx.Document()
    if narrative:
        document.add_paragraph(narrative)
    for table_rows in tables:
        n_cols = len(table_rows[0])
        table = document.add_table(rows=0, cols=n_cols)
        for row_values in table_rows:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_values):
                row_cells[i].text = value
    document.save(path)


def build_program_workbook(path: Path, rows: list[dict[int, object]]) -> None:
    """rows: list of sparse {0-based column index: value} row specs, matching how the
    real importer indexes row tuples (row[0], row[22], row[55], etc.)."""
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Selected Programs"
    for row_index, row_data in enumerate(rows, start=1):
        for col_index, value in row_data.items():
            worksheet.cell(row=row_index, column=col_index + 1, value=value)
    workbook.save(path)
